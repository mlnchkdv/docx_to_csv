import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

def extract_table_from_docx(docx_file):
    doc = Document(docx_file)
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(data)
        return df  # Возвращаем первую найденную таблицу
    return None

def create_docx_from_csv(csv_file, font_name="Arial", font_size=12, orientation="Portrait"):
    df = pd.read_csv(csv_file)
    doc = Document()
    
    section = doc.sections[0]
    if orientation == "Landscape":
        section.orientation = 1  # 1 соответствует альбомной ориентации
    
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = col_name
        run = table.cell(0, j).paragraphs[0].runs[0]
        run.font.name = font_name
        run.font.size = font_size
    
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            run = cell.paragraphs[0].runs[0]
            run.font.name = font_name
            run.font.size = font_size
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("DOCX <-> CSV Converter")
    
    uploaded_docx = st.file_uploader("Upload a DOCX file", type=["docx"])
    if uploaded_docx:
        df = extract_table_from_docx(uploaded_docx)
        if df is not None:
            st.write("Extracted Table:")
            st.dataframe(df)
            column_names = [st.text_input(f"Column {i+1} name", value=f"Column_{i+1}") for i in range(df.shape[1])]
            df.columns = column_names
            
            buffer = BytesIO()
            df.to_csv(buffer, index=False, encoding='utf-8')
            buffer.seek(0)
            
            st.download_button("Download CSV", data=buffer, file_name="converted_table.csv", mime="text/csv")
        else:
            st.error("No tables found in the document.")
    
    uploaded_csv = st.file_uploader("Upload a CSV file", type=["csv"])
    if uploaded_csv:
        font_name = st.selectbox("Choose font", ["Arial", "Times New Roman", "Courier New"])
        font_size = st.slider("Select font size", 8, 24, 12)
        orientation = st.selectbox("Page orientation", ["Portrait", "Landscape"])
        
        docx_buffer = create_docx_from_csv(uploaded_csv, font_name, font_size, orientation)
        
        st.download_button("Download DOCX", data=docx_buffer, file_name="converted_table.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
